import os
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import pytz
from docx.enum.section import WD_SECTION
import smtplib
from email.message import EmailMessage

# === Paths ===
base_dir = "/home/nipun/vapt_project/vapt_reports"
summary_dir = os.path.join(base_dir, "summary")
report_output_dir = "/home/nipun/vapt_project/vapt_reports/reports"
os.makedirs(report_output_dir, exist_ok=True)

# === Create Word Document ===
doc = Document()
section = doc.sections[0]

# === COVER PAGE ===
doc.add_paragraph()

# Title
cover_title = doc.add_paragraph("Network Resiliency Report")
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_run = cover_title.runs[0]
cover_run.font.size = Pt(28)
cover_run.font.bold = True

# Subtitle
cover_subtitle = doc.add_paragraph("Automated Audit & Compliance Summary")
cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_run2 = cover_subtitle.runs[0]
cover_run2.font.size = Pt(18)
cover_run2.italic = True

# Date
cover_date = doc.add_paragraph(datetime.now().strftime("Generated on: %B %d, %Y"))
cover_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_date.runs[0].italic = True

# Organization
doc.add_paragraph("\n\n")
org_para = doc.add_paragraph("Nipun Net Solutions Pvt. Ltd.")
org_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
org_para.runs[0].font.size = Pt(14)

# Add Logo to Cover Page (Centered)
logo_name = [f for f in os.listdir(os.path.dirname(__file__)) if f.startswith("Nipun_Logo") and f.endswith(".jpeg")]
if not logo_name:
    raise FileNotFoundError("Logo file not found in script directory.")
logo_path = os.path.join(os.path.dirname(__file__), logo_name[0])
logo_para = doc.add_paragraph()
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
logo_para_run = logo_para.add_run()
logo_para_run.add_picture(logo_path, width=Inches(2.0))

# Page break after cover page
doc.add_page_break()

# === HEADER ===
header = section.header
header_table = header.add_table(rows=2, cols=1, width=Inches(6))
header_table.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# Row 1: Logo
logo_row = header_table.rows[0].cells[0].paragraphs[0]
logo_row.alignment = WD_ALIGN_PARAGRAPH.RIGHT
logo_run = logo_row.add_run()
logo_run.add_picture(logo_path, width=Inches(1.0))

# Row 2: Line below header
line_row = header_table.rows[1].cells[0].paragraphs[0]
p_border = line_row._p.get_or_add_pPr()
border = OxmlElement("w:pBdr")
bottom = OxmlElement("w:bottom")
bottom.set(qn("w:val"), "single")
bottom.set(qn("w:sz"), "6")
bottom.set(qn("w:space"), "1")
bottom.set(qn("w:color"), "auto")
border.append(bottom)
p_border.append(border)

# === FOOTER ===
footer = section.footer
footer_table = footer.add_table(rows=2, cols=2, width=Inches(6))
footer_table.autofit = True
footer_table.allow_autofit = True

# Row 1: Line above footer
line_footer = footer_table.rows[0].cells[0].merge(footer_table.rows[0].cells[1])
line_para = line_footer.paragraphs[0]
p_border_footer = line_para._p.get_or_add_pPr()
top = OxmlElement("w:top")
top.set(qn("w:val"), "single")
top.set(qn("w:sz"), "6")
top.set(qn("w:space"), "1")
top.set(qn("w:color"), "auto")
border_footer = OxmlElement("w:pBdr")
border_footer.append(top)
p_border_footer.append(border_footer)

# Row 2: Left - Nipun Confidential, Right - Page x of y
left_cell = footer_table.rows[1].cells[0]
right_cell = footer_table.rows[1].cells[1]

# Left side
left_para = left_cell.paragraphs[0]
left_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run_left = left_para.add_run("Nipun Confidential")
run_left.font.name = "Arial"
run_left.font.size = Pt(10)
run_left.font.bold = True

# Right side
right_para = right_cell.paragraphs[0]
right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run_right = right_para.add_run()
run_right.font.name = "Arial"
run_right.font.size = Pt(10)
run_right.font.bold = True

# Add PAGE x OF y field
for instr in ["PAGE", "NUMPAGES"]:
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")

    instrText = OxmlElement("w:instrText")
    instrText.text = instr

    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")

    fldChar3 = OxmlElement("w:fldChar")
    fldChar3.set(qn("w:fldCharType"), "end")

    run_right._r.append(fldChar1)
    run_right._r.append(instrText)
    run_right._r.append(fldChar2)
    run_right._r.append(fldChar3)

    if instr == "PAGE":
        run_right.add_text(" of ")

# === Title ===

# === Summary Table ===
latest_summary_file = sorted([f for f in os.listdir(summary_dir) if f.endswith(".xlsx")])[-1]
summary_path = os.path.join(summary_dir, latest_summary_file)
summary_df = pd.read_excel(summary_path)

doc.add_heading("Overall Device Compliance Summary", level=1)
table = doc.add_table(rows=1, cols=len(summary_df.columns))
table.style = "Table Grid"
hdr_cells = table.rows[0].cells
for i, column in enumerate(summary_df.columns):
    hdr_cells[i].text = str(column)

for _, row in summary_df.iterrows():
    row_cells = table.add_row().cells
    for i, cell in enumerate(row):
        row_cells[i].text = str(cell)

# === Per-device reports ===
for device in os.listdir(base_dir):
    device_path = os.path.join(base_dir, device)
    if not os.path.isdir(device_path) or device in ("summary", "reports"):
        continue

    timestamps = sorted([t for t in os.listdir(device_path) if os.path.isdir(os.path.join(device_path, t))])
    if not timestamps:
        continue

    latest_ts = timestamps[-1]
    report_dir = os.path.join(device_path, latest_ts)

    doc.add_page_break()
    doc.add_heading(f"Device: {device}", level=1)

    # Hardening Summary
    summary_file = os.path.join(report_dir, "hardening_summary.xlsx")
    if os.path.exists(summary_file):
        doc.add_heading("Hardening Summary", level=2)
        try:
            df = pd.read_excel(summary_file, sheet_name="Hardening Report")
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            for i, column in enumerate(df.columns):
                hdr_cells[i].text = str(column)
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    row_cells[i].text = str(cell)
        except Exception as e:
            doc.add_paragraph(f"Error reading hardening summary: {e}")

    # Running Config
    config_file = os.path.join(report_dir, "running_config.txt")
    if os.path.exists(config_file):
        doc.add_heading("Running Configuration", level=2)
        with open(config_file, "r") as f:
            config_text = f.read()
            doc.add_paragraph(config_text)

    # Nmap Scan
    nmap_file = os.path.join(report_dir, "nmap.txt")
    if os.path.exists(nmap_file):
        doc.add_heading("Nmap Scan Results", level=2)
        with open(nmap_file, "r") as f:
            nmap_text = f.read()
            doc.add_paragraph(nmap_text)

# === Save Document ===
ist = pytz.timezone("Asia/Kolkata")
timestamp = datetime.now(ist).strftime("%Y%m%d_%H%M%S")
output_path = os.path.join(report_output_dir, f"Final_Report_{timestamp}.docx")
doc.save(output_path)
print(f"✅ Word report saved at: {output_path}")

# Email Configuration
EMAIL_SENDER = "manojpassive1234@gmail.com"
EMAIL_PASSWORD = "apjyuuwsiyyblctq"
EMAIL_RECEIVER = "manoj.j@nipun.net"
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 587

# Find the latest Word report
latest_report = sorted(
    [f for f in os.listdir(report_output_dir) if f.endswith(".docx")],
    key=lambda x: os.path.getmtime(os.path.join(report_output_dir, x)),
    reverse=True
)[0]
report_path = os.path.join(report_output_dir, latest_report)

# Create Email Message
msg = EmailMessage()
msg["Subject"] = "Network Resiliency Final Report"
msg["From"] = EMAIL_SENDER
msg["To"] = EMAIL_RECEIVER
msg.set_content("Hi team,\n\nPlease find the attached Final Network Resiliency Report.\n\nRegards,\nNetwork Automation Team,\nNipun Net Solutions PVT LTD")

# Attach Word Document
with open(report_path, "rb") as file:
    file_data = file.read()
    msg.add_attachment(file_data, maintype="application", subtype="vnd.openxmlformats-officedocument.wordprocessingml.document", filename=latest_report)

# Send Email
try:
    with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as smtp:
        smtp.starttls()
        smtp.login(EMAIL_SENDER, EMAIL_PASSWORD)
        smtp.send_message(msg)
    print("✅ Email sent successfully.")
except Exception as e:
    print(f"❌ Failed to send email: {e}")
